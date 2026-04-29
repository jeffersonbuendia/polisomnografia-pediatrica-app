import streamlit as st
from datetime import date
from io import BytesIO
from html import escape

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    KeepTogether,
)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

st.set_page_config(
    page_title="PSG Pediátrica - Reporte hospitalario",
    page_icon="💤",
    layout="wide"
)

st.markdown(
    """
    <style>
    .main-title {
        font-size: 2.35rem;
        font-weight: 850;
        margin-bottom: 0.2rem;
        color: #111827;
    }
    .subtitle {
        color: #4B5563;
        font-size: 1.02rem;
        margin-bottom: 1.2rem;
    }
    .section-card {
        background-color: #F9FAFB;
        border: 1px solid #E5E7EB;
        border-radius: 16px;
        padding: 1.15rem;
        margin-bottom: 1rem;
    }
    .report-box {
        background-color: #FFFFFF;
        border: 1px solid #D1D5DB;
        border-radius: 16px;
        padding: 1.2rem;
        line-height: 1.65;
        box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    }
    .small-muted {
        color: #6B7280;
        font-size: 0.88rem;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown('<div class="main-title">Calculadora clínica de polisomnografía pediátrica</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="subtitle">Versión hospitalaria premium: genera reporte clínico estructurado en PDF y Word editable.</div>',
    unsafe_allow_html=True
)


# ============================================================
# FUNCIONES CLÍNICAS
# ============================================================

def clasificar_iah(iah: float) -> str:
    if iah < 1:
        return "sin evidencia polisomnográfica de apnea obstructiva del sueño pediátrica"
    if iah < 5:
        return "apnea obstructiva del sueño pediátrica leve"
    if iah < 10:
        return "apnea obstructiva del sueño pediátrica moderada"
    return "apnea obstructiva del sueño pediátrica severa"


def grado_corto_iah(iah: float) -> str:
    if iah < 1:
        return "normal"
    if iah < 5:
        return "leve"
    if iah < 10:
        return "moderada"
    return "severa"


def interpretar_oxigenacion(spo2_min: float, tiempo_menor_90: float) -> str:
    if spo2_min < 80 or tiempo_menor_90 > 5:
        return "compromiso significativo de la oxigenación nocturna"
    if spo2_min < 90 or tiempo_menor_90 > 0:
        return "desaturación nocturna leve a moderada"
    return "oxigenación nocturna conservada"


def interpretar_eficiencia(eficiencia: float) -> str:
    if eficiencia >= 85:
        return "eficiencia del sueño conservada"
    if eficiencia >= 75:
        return "eficiencia del sueño discretamente reducida"
    return "eficiencia del sueño reducida"


def interpretar_arousals(indice_arousal: float) -> str:
    if indice_arousal < 10:
        return "sin fragmentación significativa del sueño"
    if indice_arousal < 20:
        return "fragmentación leve a moderada del sueño"
    return "fragmentación importante del sueño"


def interpretar_plmi(plmi: float) -> str:
    if plmi < 5:
        return "sin incremento patológico de movimientos periódicos de extremidades"
    return "incremento del índice de movimientos periódicos de extremidades"


def patron_rem(iah_rem: float, iah_nrem: float) -> str:
    if iah_nrem == 0 and iah_rem > 0:
        return "con predominio durante sueño REM"
    if iah_rem >= iah_nrem * 1.5:
        return "con predominio durante sueño REM"
    return "sin predominio REM claramente dominante"


def patron_posicional(iah_supino: float, iah_lateral: float) -> str:
    if iah_lateral == 0 and iah_supino > 0:
        return "con componente posicional en decúbito supino"
    if iah_supino >= iah_lateral * 2:
        return "con componente posicional en decúbito supino"
    return "sin patrón posicional claramente dominante"


def safe_text(value) -> str:
    return escape(str(value)).replace("\n", "<br/>")


def make_impression(data: dict) -> str:
    return (
        f"La polisomnografía pediátrica muestra un tiempo total de sueño de "
        f"{data['tiempo_sueno']:.1f} minutos y una eficiencia del sueño de "
        f"{data['eficiencia']:.1f}%, compatible con {data['interp_eficiencia']}. "
        f"El índice de apnea-hipopnea fue de {data['iah']:.1f} eventos/hora, "
        f"con componente obstructivo de {data['iah_obstructivo']:.1f} eventos/hora, "
        f"hallazgo que clasifica el estudio como {data['diagnostico_iah']}. "
        f"Desde el punto de vista fisiológico, la saturación basal fue de "
        f"{data['spo2_basal']:.1f}%, con nadir de {data['spo2_min']:.1f}% y "
        f"tiempo con SpO₂ <90% de {data['tiempo_menor_90']:.1f}% del tiempo total de sueño, "
        f"lo cual sugiere {data['interp_oxigenacion']}. El índice de arousal fue de "
        f"{data['indice_arousal']:.1f} eventos/hora, indicando {data['interp_arousal']}. "
        f"El análisis por fases mostró IAH REM de {data['iah_rem']:.1f} eventos/hora e "
        f"IAH NREM de {data['iah_nrem']:.1f} eventos/hora, por lo que el patrón se interpreta "
        f"como {data['frase_rem']}. El análisis posicional muestra {data['frase_posicion']}. "
        f"En conjunto, los hallazgos son consistentes con trastorno respiratorio obstructivo "
        f"del sueño pediátrico de severidad {data['grado_corto']}, con repercusión sobre la "
        f"oxigenación descrita como {data['interp_oxigenacion']} y fragmentación del sueño "
        f"descrita como {data['interp_arousal']}."
    )


# ============================================================
# PDF
# ============================================================

def build_pdf(data: dict) -> bytes:
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=1.45 * cm,
        leftMargin=1.45 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.35 * cm,
        title="Reporte de polisomnografía pediátrica",
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#111827"),
        spaceAfter=6,
    )

    institution_style = ParagraphStyle(
        "InstitutionCustom",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1F2937"),
        spaceAfter=2,
    )

    date_style = ParagraphStyle(
        "DateCustom",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=10,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#4B5563"),
        spaceAfter=8,
    )

    h_style = ParagraphStyle(
        "HeadingCustom",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=colors.HexColor("#111827"),
        spaceBefore=7,
        spaceAfter=4,
    )

    normal_style = ParagraphStyle(
        "NormalCustom",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.8,
        leading=11.5,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor("#111827"),
    )

    note_style = ParagraphStyle(
        "NoteCustom",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=7.8,
        leading=9.6,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#4B5563"),
    )

    story = []

    if data.get("institucion", "").strip():
        story.append(Paragraph(safe_text(data["institucion"]), institution_style))

    story.append(Paragraph("REPORTE DE POLISOMNOGRAFÍA PEDIÁTRICA", title_style))
    story.append(Paragraph(f"Fecha del estudio: {data['fecha_estudio']}", date_style))

    def section(title):
        story.append(Paragraph(title, h_style))

    def p(text):
        story.append(Paragraph(text, normal_style))
        story.append(Spacer(1, 4))

    def table(rows, widths=None):
        table_obj = Table(rows, colWidths=widths, hAlign="LEFT")
        table_obj.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E5E7EB")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111827")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.2),
            ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
            ("ALIGN", (0, 0), (0, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table_obj)
        story.append(Spacer(1, 5))

    section("1. Indicación clínica")
    p(safe_text(data["indicacion"]))

    section("2. Condiciones del estudio")
    table([
        ["Parámetro", "Resultado"],
        ["Tipo de estudio", data["tipo_estudio"]],
        ["Duración total del registro", f'{data["duracion_registro"]:.1f} min'],
        ["Tiempo en cama (TIB)", f'{data["tiempo_cama"]:.1f} min'],
        ["Tiempo total de sueño (TST)", f'{data["tiempo_sueno"]:.1f} min'],
        ["Eficiencia del sueño", f'{data["eficiencia"]:.1f}%'],
        ["Latencia al sueño", f'{data["latencia_sueno"]:.1f} min'],
        ["Latencia REM", f'{data["latencia_rem"]:.1f} min'],
        ["Posición predominante", data["posicion_predominante"]],
    ], widths=[9 * cm, 7.3 * cm])

    section("3. Arquitectura del sueño")
    table([
        ["Estadio", "% del TST"],
        ["N1", f'{data["n1"]:.1f}%'],
        ["N2", f'{data["n2"]:.1f}%'],
        ["N3", f'{data["n3"]:.1f}%'],
        ["REM", f'{data["rem"]:.1f}%'],
    ], widths=[9 * cm, 7.3 * cm])
    p(f'Interpretación: el estudio muestra {data["interp_eficiencia"]}. La arquitectura debe analizarse con la edad, la duración del sueño REM y el grado de fragmentación.')

    section("4. Eventos respiratorios")
    table([
        ["Parámetro", "Resultado"],
        ["IAH total", f'{data["iah"]:.1f} eventos/hora'],
        ["IAH obstructivo", f'{data["iah_obstructivo"]:.1f} eventos/hora'],
        ["IAH central", f'{data["iah_central"]:.1f} eventos/hora'],
        ["Índice de hipopneas", f'{data["indice_hipopneas"]:.1f} eventos/hora'],
        ["IAH REM", f'{data["iah_rem"]:.1f} eventos/hora'],
        ["IAH NREM", f'{data["iah_nrem"]:.1f} eventos/hora'],
    ], widths=[9 * cm, 7.3 * cm])
    p(f'Interpretación: hallazgos compatibles con {data["diagnostico_iah"]}, {data["frase_rem"]}.')

    section("5. Oxigenación")
    table([
        ["Parámetro", "Resultado"],
        ["SpO₂ basal", f'{data["spo2_basal"]:.1f}%'],
        ["SpO₂ mínima", f'{data["spo2_min"]:.1f}%'],
        ["Tiempo con SpO₂ <90%", f'{data["tiempo_menor_90"]:.1f}% del TST'],
        ["ODI", f'{data["odi"]:.1f} eventos/hora'],
    ], widths=[9 * cm, 7.3 * cm])
    p(f'Interpretación: se documenta {data["interp_oxigenacion"]}.')

    section("6. Eventos cardiovasculares")
    table([
        ["Parámetro", "Resultado"],
        ["Frecuencia cardíaca promedio", f'{data["fc_promedio"]:.1f} lpm'],
        ["Frecuencia cardíaca mínima", f'{data["fc_min"]:.1f} lpm'],
        ["Frecuencia cardíaca máxima", f'{data["fc_max"]:.1f} lpm'],
        ["Arritmias", data["arritmias"]],
    ], widths=[9 * cm, 7.3 * cm])

    section("7. Movimientos periódicos de extremidades")
    table([
        ["Parámetro", "Resultado"],
        ["PLMI", f'{data["plmi"]:.1f} eventos/hora'],
    ], widths=[9 * cm, 7.3 * cm])
    p(f'Interpretación: {data["interp_plmi"]}.')

    section("8. Microdespertares")
    table([
        ["Parámetro", "Resultado"],
        ["Índice de arousal", f'{data["indice_arousal"]:.1f} eventos/hora'],
    ], widths=[9 * cm, 7.3 * cm])
    p(f'Interpretación: {data["interp_arousal"]}.')

    section("9. Análisis posicional")
    table([
        ["Posición", "IAH"],
        ["Supino", f'{data["iah_supino"]:.1f} eventos/hora'],
        ["Lateral", f'{data["iah_lateral"]:.1f} eventos/hora'],
    ], widths=[9 * cm, 7.3 * cm])
    p(f'Interpretación: {data["frase_posicion"]}.')

    section("10. Calidad técnica del estudio")
    p(f'Calidad técnica: <b>{safe_text(data["calidad"])}</b>.<br/>Observaciones técnicas: {safe_text(data["artefactos"])}')

    section("11. Conclusión diagnóstica")
    p(
        f'Estudio compatible con <b>{data["diagnostico_iah"]}</b>, con IAH total de '
        f'<b>{data["iah"]:.1f} eventos/hora</b>, IAH obstructivo de '
        f'<b>{data["iah_obstructivo"]:.1f} eventos/hora</b>, saturación mínima de '
        f'<b>{data["spo2_min"]:.1f}%</b> y {data["interp_oxigenacion"]}. '
        f'El patrón respiratorio fue descrito como {data["frase_rem"]} y {data["frase_posicion"]}.'
    )

    section("12. Impresión clínica integrada")
    p(safe_text(data["impresion"]))

    if data.get("mostrar_nota"):
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            "Nota técnica: la clasificación automática usa umbrales pediátricos habituales del IAH "
            "(normal <1; leve 1 a <5; moderada 5 a <10; severa >=10 eventos/hora). "
            "La interpretación final debe integrarse con edad, síntomas, comorbilidades, examen físico "
            "y criterios vigentes del laboratorio de sueño.",
            note_style
        ))

    if data.get("medico_lector", "").strip() or data.get("registro_medico", "").strip():
        story.append(Spacer(1, 12))
        firma = [
            ["Médico lector", data.get("medico_lector", "")],
            ["Registro profesional", data.get("registro_medico", "")],
        ]
        table([["Firma / responsable", ""]] + firma, widths=[6 * cm, 10.3 * cm])

    doc.build(story)
    return buffer.getvalue()


# ============================================================
# WORD
# ============================================================

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], rows[0][0], bold=True)
    set_cell_text(hdr[1], rows[0][1], bold=True)
    set_cell_shading(hdr[0], "E5E7EB")
    set_cell_shading(hdr[1], "E5E7EB")

    for left, right in rows[1:]:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)
    doc.add_paragraph("")


def add_docx_section(doc, title):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(17, 24, 39)


def add_docx_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)


def build_docx(data: dict) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    if data.get("institucion", "").strip():
        p_inst = doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_inst = p_inst.add_run(data["institucion"])
        r_inst.bold = True
        r_inst.font.size = Pt(11)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("REPORTE DE POLISOMNOGRAFÍA PEDIÁTRICA")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(17, 24, 39)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run(f"Fecha del estudio: {data['fecha_estudio']}")
    r_date.font.size = Pt(9)
    r_date.font.color.rgb = RGBColor(75, 85, 99)

    add_docx_section(doc, "1. Indicación clínica")
    add_docx_paragraph(doc, data["indicacion"])

    add_docx_section(doc, "2. Condiciones del estudio")
    add_docx_table(doc, [
        ["Parámetro", "Resultado"],
        ["Tipo de estudio", data["tipo_estudio"]],
        ["Duración total del registro", f'{data["duracion_registro"]:.1f} min'],
        ["Tiempo en cama (TIB)", f'{data["tiempo_cama"]:.1f} min'],
        ["Tiempo total de sueño (TST)", f'{data["tiempo_sueno"]:.1f} min'],
        ["Eficiencia del sueño", f'{data["eficiencia"]:.1f}%'],
        ["Latencia al sueño", f'{data["latencia_sueno"]:.1f} min'],
        ["Latencia REM", f'{data["latencia_rem"]:.1f} min'],
        ["Posición predominante", data["posicion_predominante"]],
    ])

    add_docx_section(doc, "3. Arquitectura del sueño")
    add_docx_table(doc, [
        ["Estadio", "% del TST"],
        ["N1", f'{data["n1"]:.1f}%'],
        ["N2", f'{data["n2"]:.1f}%'],
        ["N3", f'{data["n3"]:.1f}%'],
        ["REM", f'{data["rem"]:.1f}%'],
    ])
    add_docx_paragraph(doc, f'Interpretación: el estudio muestra {data["interp_eficiencia"]}. La arquitectura debe analizarse con la edad, la duración del sueño REM y el grado de fragmentación.')

    add_docx_section(doc, "4. Eventos respiratorios")
    add_docx_table(doc, [
        ["Parámetro", "Resultado"],
        ["IAH total", f'{data["iah"]:.1f} eventos/hora'],
        ["IAH obstructivo", f'{data["iah_obstructivo"]:.1f} eventos/hora'],
        ["IAH central", f'{data["iah_central"]:.1f} eventos/hora'],
        ["Índice de hipopneas", f'{data["indice_hipopneas"]:.1f} eventos/hora'],
        ["IAH REM", f'{data["iah_rem"]:.1f} eventos/hora'],
        ["IAH NREM", f'{data["iah_nrem"]:.1f} eventos/hora'],
    ])
    add_docx_paragraph(doc, f'Interpretación: hallazgos compatibles con {data["diagnostico_iah"]}, {data["frase_rem"]}.')

    add_docx_section(doc, "5. Oxigenación")
    add_docx_table(doc, [
        ["Parámetro", "Resultado"],
        ["SpO₂ basal", f'{data["spo2_basal"]:.1f}%'],
        ["SpO₂ mínima", f'{data["spo2_min"]:.1f}%'],
        ["Tiempo con SpO₂ <90%", f'{data["tiempo_menor_90"]:.1f}% del TST'],
        ["ODI", f'{data["odi"]:.1f} eventos/hora'],
    ])
    add_docx_paragraph(doc, f'Interpretación: se documenta {data["interp_oxigenacion"]}.')

    add_docx_section(doc, "6. Eventos cardiovasculares")
    add_docx_table(doc, [
        ["Parámetro", "Resultado"],
        ["Frecuencia cardíaca promedio", f'{data["fc_promedio"]:.1f} lpm'],
        ["Frecuencia cardíaca mínima", f'{data["fc_min"]:.1f} lpm'],
        ["Frecuencia cardíaca máxima", f'{data["fc_max"]:.1f} lpm'],
        ["Arritmias", data["arritmias"]],
    ])

    add_docx_section(doc, "7. Movimientos periódicos de extremidades")
    add_docx_table(doc, [
        ["Parámetro", "Resultado"],
        ["PLMI", f'{data["plmi"]:.1f} eventos/hora'],
    ])
    add_docx_paragraph(doc, f'Interpretación: {data["interp_plmi"]}.')

    add_docx_section(doc, "8. Microdespertares")
    add_docx_table(doc, [
        ["Parámetro", "Resultado"],
        ["Índice de arousal", f'{data["indice_arousal"]:.1f} eventos/hora'],
    ])
    add_docx_paragraph(doc, f'Interpretación: {data["interp_arousal"]}.')

    add_docx_section(doc, "9. Análisis posicional")
    add_docx_table(doc, [
        ["Posición", "IAH"],
        ["Supino", f'{data["iah_supino"]:.1f} eventos/hora'],
        ["Lateral", f'{data["iah_lateral"]:.1f} eventos/hora'],
    ])
    add_docx_paragraph(doc, f'Interpretación: {data["frase_posicion"]}.')

    add_docx_section(doc, "10. Calidad técnica del estudio")
    add_docx_paragraph(doc, f'Calidad técnica: {data["calidad"]}. Observaciones técnicas: {data["artefactos"]}')

    add_docx_section(doc, "11. Conclusión diagnóstica")
    add_docx_paragraph(
        doc,
        f'Estudio compatible con {data["diagnostico_iah"]}, con IAH total de '
        f'{data["iah"]:.1f} eventos/hora, IAH obstructivo de '
        f'{data["iah_obstructivo"]:.1f} eventos/hora, saturación mínima de '
        f'{data["spo2_min"]:.1f}% y {data["interp_oxigenacion"]}. '
        f'El patrón respiratorio fue descrito como {data["frase_rem"]} y {data["frase_posicion"]}.'
    )

    add_docx_section(doc, "12. Impresión clínica integrada")
    add_docx_paragraph(doc, data["impresion"])

    if data.get("mostrar_nota"):
        add_docx_section(doc, "Nota técnica")
        p = doc.add_paragraph()
        r = p.add_run(
            "La clasificación automática usa umbrales pediátricos habituales del IAH "
            "(normal <1; leve 1 a <5; moderada 5 a <10; severa >=10 eventos/hora). "
            "La interpretación final debe integrarse con edad, síntomas, comorbilidades, examen físico "
            "y criterios vigentes del laboratorio de sueño."
        )
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(75, 85, 99)

    if data.get("medico_lector", "").strip() or data.get("registro_medico", "").strip():
        doc.add_paragraph("")
        add_docx_table(doc, [
            ["Firma / responsable", ""],
            ["Médico lector", data.get("medico_lector", "")],
            ["Registro profesional", data.get("registro_medico", "")],
        ])

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ============================================================
# FORMULARIO STREAMLIT
# ============================================================

with st.sidebar:
    st.header("Encabezado institucional")
    institucion = st.text_input("Institución / laboratorio", value="")
    medico_lector = st.text_input("Médico lector", value="")
    registro_medico = st.text_input("Registro profesional", value="")
    mostrar_nota = st.checkbox("Incluir nota técnica", value=True)

st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.header("1. Indicación clínica")
indicacion = st.text_area(
    "Indicación",
    value="Evaluación de trastorno respiratorio del sueño en paciente pediátrico.",
    height=80
)
st.markdown('</div>', unsafe_allow_html=True)

st.header("2. Condiciones del estudio")
col1, col2, col3 = st.columns(3)

with col1:
    fecha_estudio = st.date_input("Fecha del estudio", value=date.today())
    tipo_estudio = st.selectbox(
        "Tipo de estudio",
        [
            "Polisomnografía nocturna completa nivel I",
            "Polisomnografía hospitalaria",
            "Polisomnografía ambulatoria",
            "Otro"
        ]
    )
    duracion_registro = st.number_input("Duración total del registro (min)", min_value=0.0, value=480.0, step=5.0)

with col2:
    tiempo_cama = st.number_input("Tiempo en cama - TIB (min)", min_value=0.0, value=480.0, step=5.0)
    tiempo_sueno = st.number_input("Tiempo total de sueño - TST (min)", min_value=0.0, value=420.0, step=5.0)
    eficiencia = st.number_input("Eficiencia del sueño (%)", min_value=0.0, max_value=100.0, value=87.5, step=0.1)

with col3:
    latencia_sueno = st.number_input("Latencia al sueño (min)", min_value=0.0, value=15.0, step=1.0)
    latencia_rem = st.number_input("Latencia REM (min)", min_value=0.0, value=90.0, step=1.0)
    posicion_predominante = st.selectbox("Posición predominante", ["Supino", "Lateral", "Prono", "Mixta"])

st.header("3. Arquitectura del sueño")
col1, col2, col3, col4 = st.columns(4)

with col1:
    n1 = st.number_input("N1 (% TST)", min_value=0.0, max_value=100.0, value=5.0, step=0.1)
with col2:
    n2 = st.number_input("N2 (% TST)", min_value=0.0, max_value=100.0, value=45.0, step=0.1)
with col3:
    n3 = st.number_input("N3 (% TST)", min_value=0.0, max_value=100.0, value=25.0, step=0.1)
with col4:
    rem = st.number_input("REM (% TST)", min_value=0.0, max_value=100.0, value=25.0, step=0.1)

suma_estadios = n1 + n2 + n3 + rem
if abs(suma_estadios - 100) > 2:
    st.warning(f"La suma de estadios es {suma_estadios:.1f}%. Revise si los porcentajes están completos.")

st.header("4. Eventos respiratorios")
col1, col2, col3, col4 = st.columns(4)

with col1:
    iah = st.number_input("IAH total (eventos/hora)", min_value=0.0, value=3.2, step=0.1)
with col2:
    iah_obstructivo = st.number_input("IAH obstructivo (eventos/hora)", min_value=0.0, value=3.0, step=0.1)
with col3:
    iah_central = st.number_input("IAH central (eventos/hora)", min_value=0.0, value=0.2, step=0.1)
with col4:
    indice_hipopneas = st.number_input("Índice de hipopneas (eventos/hora)", min_value=0.0, value=2.5, step=0.1)

col1, col2 = st.columns(2)
with col1:
    iah_rem = st.number_input("IAH en REM (eventos/hora)", min_value=0.0, value=6.0, step=0.1)
with col2:
    iah_nrem = st.number_input("IAH en NREM (eventos/hora)", min_value=0.0, value=2.0, step=0.1)

st.header("5. Oxigenación")
col1, col2, col3, col4 = st.columns(4)

with col1:
    spo2_basal = st.number_input("SpO₂ basal (%)", min_value=0.0, max_value=100.0, value=96.0, step=0.1)
with col2:
    spo2_min = st.number_input("SpO₂ mínima (%)", min_value=0.0, max_value=100.0, value=88.0, step=0.1)
with col3:
    tiempo_menor_90 = st.number_input("Tiempo con SpO₂ <90% (% TST)", min_value=0.0, max_value=100.0, value=1.5, step=0.1)
with col4:
    odi = st.number_input("ODI / índice de desaturación", min_value=0.0, value=4.0, step=0.1)

st.header("6. Eventos cardiovasculares y movimientos")
col1, col2, col3, col4 = st.columns(4)

with col1:
    fc_promedio = st.number_input("Frecuencia cardíaca promedio", min_value=0.0, value=85.0, step=1.0)
with col2:
    fc_min = st.number_input("Frecuencia cardíaca mínima", min_value=0.0, value=62.0, step=1.0)
with col3:
    fc_max = st.number_input("Frecuencia cardíaca máxima", min_value=0.0, value=125.0, step=1.0)
with col4:
    arritmias = st.selectbox("Arritmias", ["No", "Sí"])

col1, col2 = st.columns(2)
with col1:
    plmi = st.number_input("PLMI (eventos/hora)", min_value=0.0, value=2.0, step=0.1)
with col2:
    indice_arousal = st.number_input("Índice de arousal (eventos/hora)", min_value=0.0, value=8.0, step=0.1)

st.header("7. Análisis posicional")
col1, col2 = st.columns(2)

with col1:
    iah_supino = st.number_input("IAH en supino", min_value=0.0, value=5.0, step=0.1)
with col2:
    iah_lateral = st.number_input("IAH en lateral", min_value=0.0, value=2.0, step=0.1)

st.header("8. Calidad técnica del estudio")
calidad = st.selectbox("Calidad técnica", ["Adecuada", "Limitada", "Subóptima"])
artefactos = st.text_area(
    "Canales perdidos o artefactos relevantes",
    value="Sin artefactos significativos que limiten la interpretación.",
    height=80
)

# ============================================================
# CÁLCULOS DERIVADOS Y SALIDA
# ============================================================

diagnostico_iah = clasificar_iah(iah)
grado_corto = grado_corto_iah(iah)
interp_oxigenacion = interpretar_oxigenacion(spo2_min, tiempo_menor_90)
interp_eficiencia = interpretar_eficiencia(eficiencia)
interp_arousal = interpretar_arousals(indice_arousal)
interp_plmi = interpretar_plmi(plmi)
frase_rem = patron_rem(iah_rem, iah_nrem)
frase_posicion = patron_posicional(iah_supino, iah_lateral)

data = {
    "institucion": institucion,
    "medico_lector": medico_lector,
    "registro_medico": registro_medico,
    "mostrar_nota": mostrar_nota,
    "indicacion": indicacion,
    "fecha_estudio": fecha_estudio,
    "tipo_estudio": tipo_estudio,
    "duracion_registro": duracion_registro,
    "tiempo_cama": tiempo_cama,
    "tiempo_sueno": tiempo_sueno,
    "eficiencia": eficiencia,
    "latencia_sueno": latencia_sueno,
    "latencia_rem": latencia_rem,
    "posicion_predominante": posicion_predominante,
    "n1": n1,
    "n2": n2,
    "n3": n3,
    "rem": rem,
    "iah": iah,
    "iah_obstructivo": iah_obstructivo,
    "iah_central": iah_central,
    "indice_hipopneas": indice_hipopneas,
    "iah_rem": iah_rem,
    "iah_nrem": iah_nrem,
    "spo2_basal": spo2_basal,
    "spo2_min": spo2_min,
    "tiempo_menor_90": tiempo_menor_90,
    "odi": odi,
    "fc_promedio": fc_promedio,
    "fc_min": fc_min,
    "fc_max": fc_max,
    "arritmias": arritmias,
    "plmi": plmi,
    "indice_arousal": indice_arousal,
    "iah_supino": iah_supino,
    "iah_lateral": iah_lateral,
    "calidad": calidad,
    "artefactos": artefactos,
    "diagnostico_iah": diagnostico_iah,
    "grado_corto": grado_corto,
    "interp_oxigenacion": interp_oxigenacion,
    "interp_eficiencia": interp_eficiencia,
    "interp_arousal": interp_arousal,
    "interp_plmi": interp_plmi,
    "frase_rem": frase_rem,
    "frase_posicion": frase_posicion,
}

data["impresion"] = make_impression(data)

pdf_bytes = build_pdf(data)
docx_bytes = build_docx(data)

st.divider()
st.header("Resumen automático")

m1, m2, m3, m4 = st.columns(4)
m1.metric("IAH", f"{iah:.1f}/h", grado_corto)
m2.metric("SpO₂ mínima", f"{spo2_min:.1f}%")
m3.metric("Arousal", f"{indice_arousal:.1f}/h")
m4.metric("PLMI", f"{plmi:.1f}/h")

st.markdown(
    f"""
    <div class="report-box">
    <b>Impresión clínica integrada:</b><br><br>
    {data["impresion"]}
    </div>
    """,
    unsafe_allow_html=True
)

st.markdown('<p class="small-muted">Los archivos generados no contienen identificación del paciente por diseño.</p>', unsafe_allow_html=True)

col1, col2 = st.columns(2)

with col1:
    st.download_button(
        label="📄 Descargar reporte en PDF",
        data=pdf_bytes,
        file_name="reporte_polisomnografia_pediatrica.pdf",
        mime="application/pdf",
        use_container_width=True
    )

with col2:
    st.download_button(
        label="📝 Descargar reporte en Word (.docx)",
        data=docx_bytes,
        file_name="reporte_polisomnografia_pediatrica.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
